import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()

# Title
title = doc.add_heading('SafeDrive AI - PaaS Security Audit & Deployment Guide', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('This document outlines the results of the PaaS Security Audit, the recent configuration hardening applied to the codebase, and the step-by-step guide to securely deploying the SafeDrive AI MVP to Vercel and Render.')

# Section 1: Security Audit Summary
doc.add_heading('1. Security Audit Summary', level=1)
doc.add_paragraph('The platform has a strong architectural foundation (Next.js + FastAPI + Postgres) perfect for horizontal scaling on PaaS providers. The database strictly enforces multi-tenant isolation, preventing Cross-User Data Leakage.')
doc.add_paragraph('Critical Findings Addressed:', style='List Bullet')
doc.add_paragraph('Hardcoded Secrets: Previously, docker-compose.yml contained hardcoded JWT and Postgres passwords. These have been completely extracted to environment variables (.env files).', style='List Bullet')
doc.add_paragraph('CORS Configuration: ALLOWED_ORIGINS was hardcoded. It is now dynamically injected via environment variables.', style='List Bullet')

doc.add_paragraph('Future Technical Debt (Post-MVP):', style='List Bullet')
doc.add_paragraph('JWT Storage: Migrate tokens from localStorage to HttpOnly Secure cookies to prevent XSS attacks.', style='List Bullet')
doc.add_paragraph('AI Security: Implement strict prompt sanitization and LLM output parsing to prevent prompt injection.', style='List Bullet')


# Section 2: Environment Configuration
doc.add_heading('2. Environment Configuration', level=1)
doc.add_paragraph('The codebase now relies on strict environment variable injection. Never commit `.env` files to git. Use `.env.example` as your template.')


# Section 3: Step-by-Step Deployment Guide
doc.add_heading('3. Step-by-Step Deployment Guide', level=1)
doc.add_paragraph('Follow these exact steps to securely deploy the platform to Render (Backend) and Vercel (Frontend).')

doc.add_heading('Phase A: Prepare the Backend Database (Render)', level=2)
doc.add_paragraph('1. Go to Render.com and click "New PostgreSQL".')
doc.add_paragraph('2. Name the database (e.g., safedrive-db). Choose the free or basic tier.')
doc.add_paragraph('3. Once created, copy the "Internal Database URL" (if deploying backend on Render) or "External Database URL" (if connecting from outside).')

doc.add_heading('Phase B: Deploy the Backend API (Render)', level=2)
doc.add_paragraph('1. Go to Render.com and click "New Web Service".')
doc.add_paragraph('2. Connect your GitHub repository and select the SafeDrive AI project.')
doc.add_paragraph('3. Set the "Root Directory" to `backend` (or leave empty if using docker-compose, but Render Web Services prefer direct Dockerfile selection from the backend dir).')
doc.add_paragraph('4. Environment Variables Setup (CRITICAL):', style='List Bullet')
doc.add_paragraph('Add DATABASE_URL and paste the Postgres URL from Phase A.', style='List Bullet')
doc.add_paragraph('Add SYNC_DATABASE_URL and paste the same URL (replace postgresql+asyncpg with postgresql).', style='List Bullet')
doc.add_paragraph('Add JWT_SECRET_KEY and input a long, secure random string. DO NOT use the local-dev string.', style='List Bullet')
doc.add_paragraph('Add ALLOWED_ORIGINS. Set this to your future Vercel frontend URL (e.g., https://safedrive-ai.vercel.app).', style='List Bullet')
doc.add_paragraph('Add any required LLM API keys (OPENAI_API_KEY, GEMINI_API_KEY, etc).', style='List Bullet')
doc.add_paragraph('5. Click Deploy. Wait for the service to build and start. Copy the live backend URL (e.g., https://safedrive-api.onrender.com).')

doc.add_heading('Phase C: Deploy the Frontend (Vercel)', level=2)
doc.add_paragraph('1. Go to Vercel.com and click "Add New Project".')
doc.add_paragraph('2. Import your GitHub repository.')
doc.add_paragraph('3. Set the "Root Directory" to `frontend`. Vercel will automatically detect Next.js.')
doc.add_paragraph('4. Environment Variables Setup:', style='List Bullet')
doc.add_paragraph('Add NEXT_PUBLIC_API_URL and paste your live Render backend URL from Phase B (do not include trailing slashes).', style='List Bullet')
doc.add_paragraph('5. Click Deploy. Vercel will build the frontend and provide you with a live domain.')

doc.add_heading('Phase D: Final Verification', level=2)
doc.add_paragraph('1. Ensure that the Vercel domain exactly matches the ALLOWED_ORIGINS variable in your Render backend.')
doc.add_paragraph('2. Visit your Vercel URL, attempt to log in or create an account, and verify that API requests succeed without CORS errors.')

# Save
doc.save('docs/PaaS_Security_and_Deployment_Guide.docx')
print("Document saved successfully!")
