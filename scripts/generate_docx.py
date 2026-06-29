import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()

# Title
title = doc.add_heading('Deployment Strategy for Distracted Driving Platform', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Intro
doc.add_paragraph('Based on your tech stack (Next.js frontend, Python FastAPI backend, PostgreSQL database) and your requirement to handle 100s of users reliably and cost-effectively, here are the best ways to deploy your platform.')
doc.add_paragraph('Since this is an MVP, we want to balance low cost with ease of maintenance so you can focus on building features rather than managing servers.')

# Option 1
doc.add_heading('Option 1: The "Zero-Headache" Managed Setup (Recommended for MVPs)', level=1)
doc.add_paragraph('This option uses Platform-as-a-Service (PaaS) providers. They handle the server management, SSL certificates, and scaling for you. If an app crashes, they automatically restart it.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Frontend (Next.js): ').bold = True
p.add_run('Vercel (The creators of Next.js). It provides the best performance and easiest deployment for Next.js out of the box.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Backend (Python) & Database (PostgreSQL): ').bold = True
p.add_run('Render or Railway. Both are incredibly easy to use, natively support Python and Docker, and offer managed databases with automated backups.')

doc.add_heading('Estimated Monthly Cost: ~$15 - $25 / month', level=2)
doc.add_paragraph('Vercel: $0 (The free tier is extremely generous and easily handles hundreds of users).', style='List Bullet')
doc.add_paragraph('Render/Railway Backend: ~$7 - $10 / month (Basic tier with 1GB RAM, plenty for 100s of concurrent API requests).', style='List Bullet')
doc.add_paragraph('Render/Railway Database: ~$7 - $15 / month (Managed Postgres, ensures your data is safe and backed up).', style='List Bullet')
doc.add_paragraph('Domain Name: ~$10 - $15 / year (e.g., from Namecheap or Cloudflare).', style='List Bullet')

doc.add_heading('Pros & Cons', level=2)
doc.add_paragraph('✅ Pros: Zero server maintenance, automatic GitHub deployments, extremely reliable, automatic SSL.', style='List Bullet')
doc.add_paragraph('❌ Cons: Slightly more expensive than raw servers as you scale up to thousands of users.', style='List Bullet')

# Option 2
doc.add_heading('Option 2: The "Cost-Effective" VPS Setup (Using Docker Compose)', level=1)
doc.add_paragraph('Since you already have a working docker-compose.yml file, you can rent a single Virtual Private Server (VPS), install Docker, and run everything exactly as you do on your local machine.')

doc.add_paragraph('Provider: DigitalOcean, Hetzner, or Linode.', style='List Bullet')
doc.add_paragraph("Setup: You rent a Linux server (Ubuntu), SSH into it, clone your GitHub repo, and run `docker compose up -d --build`. You will also need to set up a reverse proxy (like Nginx or Traefik) to handle domain routing and SSL certificates (via Let's Encrypt).", style='List Bullet')

doc.add_heading('Estimated Monthly Cost: ~$5 - $12 / month', level=2)
doc.add_paragraph('VPS (Virtual Server): $5 - $12 / month (A server with 2GB RAM and 1-2 vCPUs will easily handle 100s of users).', style='List Bullet')
doc.add_paragraph('Domain Name: ~$10 - $15 / year.', style='List Bullet')

doc.add_heading('Pros & Cons', level=2)
doc.add_paragraph('✅ Pros: Extremely cheap, absolute control over your environment, predictable pricing.', style='List Bullet')
doc.add_paragraph('❌ Cons: You are the sysadmin. If the server goes down, you have to fix it. You have to manually configure SSL certificates, firewall rules, and database backups.', style='List Bullet')


# Option 3
doc.add_heading('Option 3: Enterprise Cloud (AWS, Google Cloud, Azure)', level=1)
doc.add_paragraph('Using services like AWS Elastic Container Service (ECS), RDS (Relational Database Service), and CloudFront.')

doc.add_heading('Estimated Monthly Cost: ~$50 - $150+ / month', level=2)
doc.add_paragraph('AWS ECS/Fargate: ~$20 - $40 / month.', style='List Bullet')
doc.add_paragraph('AWS RDS (Database): ~$30 - $50 / month.', style='List Bullet')
doc.add_paragraph('Load Balancers & Network: ~$20 - $30 / month.', style='List Bullet')

doc.add_heading('Pros & Cons', level=2)
doc.add_paragraph('✅ Pros: Infinite scalability, enterprise-grade security and reliability.', style='List Bullet')
doc.add_paragraph('❌ Cons: Overkill for an MVP. Complex setup requiring DevOps knowledge (Terraform, IAM roles). Very easy to accidentally rack up a large bill.', style='List Bullet')


# Summary
doc.add_heading('Summary and Recommendation', level=1)
doc.add_paragraph('For your current needs (MVP, 100s of users, error-free), I strongly recommend Option 1 (Vercel + Render/Railway).')
doc.add_paragraph('While Option 2 is slightly cheaper, the $10/month you save is not worth the hours you will spend managing Linux security updates, fixing Let\'s Encrypt SSL renewals, and setting up manual database backup scripts.')
doc.add_paragraph('With Vercel + Render, your platform will be robust, updates will happen automatically when you push to GitHub, and you won\'t have to worry about server downtime.')

doc.add_heading('Next Steps if you choose Option 1:', level=2)
doc.add_paragraph('1. Push your code to a GitHub repository.')
doc.add_paragraph('2. Go to Vercel.com, connect your GitHub, and import the frontend folder.')
doc.add_paragraph('3. Go to Render.com (or Railway.app), create a new PostgreSQL database.')
doc.add_paragraph('4. Create a new Web Service on Render, connect your GitHub, point it to the backend folder (or select your Dockerfile), and add the environment variables (including the new Database URL).')

# Save
doc.save('docs/Deployment_Options.docx')
print("Document saved successfully!")
