import docx
doc = docx.Document(r'e:\Shreya Dixit Foundation\Distracted_Driving_platform\SafeDrive_AI_Calibration_Question_Bank.docx')
print("=== PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    print(repr(para.text))
print("=== TABLES ===")
for i, table in enumerate(doc.tables):
    print(f"--- TABLE {i} ---")
    for row in table.rows:
        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
        print(" | ".join(cells))
